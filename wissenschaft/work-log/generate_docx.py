#!/usr/bin/env python3
"""Generate a DOCX work-item summary from a JSON topic list.

Expects JSON on stdin with this structure:
[
  {
    "nr": 1,
    "datum": "20.02.2026",
    "umfang": "autonom, ~25\u00a0min",  # modes: dialogisch, autonom, explorativ
    "topic": "Popover-Logik überarbeitet",
    "commits": ["4b9fe0e", "c7cab06"]
  },
  ...
]

An optional tombstone entry can be included:
  {
    "type": "tombstones",
    "count": 12,
    "entries": [
      {"nr": 8, "datum": "vor dem 1.3.2026", "topic": "CORS-Fehler behoben", "commits": ["abc1234"]},
      {"nr": 9, "datum": "vor dem 15.3.2026", "topic": "Drag & Drop eingebaut", "commits": ["def5678", "ghi9012"]}
    ]
  }

Usage:
  python3 generate_docx.py <project_name> <output_path> <input_json_file>
  echo '<json>' | python3 generate_docx.py <project_name> <output_path>
"""

import json
import subprocess
import sys
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsmap
    from lxml import etree
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsmap
    from lxml import etree


def add_list_table_1_light_style(doc):
    """Inject 'List Table 1 Light' table style via XML (not in python-docx default template)."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    style_xml = f"""
    <w:style w:type="table" w:styleId="ListTable1Light"
             xmlns:w="{W}">
      <w:name w:val="List Table 1 Light"/>
      <w:basedOn w:val="TableNormal"/>
      <w:uiPriority w:val="46"/>
      <w:tblPr>
        <w:tblStyleRowBandSize w:val="1"/>
        <w:tblStyleColBandSize w:val="1"/>
      </w:tblPr>
      <w:tblStylePr w:type="firstRow">
        <w:rPr><w:b/></w:rPr>
        <w:tblPr/>
        <w:tcPr>
          <w:tcBorders>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="666666"/>
          </w:tcBorders>
        </w:tcPr>
      </w:tblStylePr>
      <w:tblStylePr w:type="lastRow">
        <w:rPr><w:b/></w:rPr>
        <w:tblPr/>
        <w:tcPr>
          <w:tcBorders>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="666666"/>
          </w:tcBorders>
        </w:tcPr>
      </w:tblStylePr>
      <w:tblStylePr w:type="band1Horz">
        <w:tblPr/>
        <w:tcPr>
          <w:shd w:val="clear" w:color="auto" w:fill="FFFFFF"/>
        </w:tcPr>
      </w:tblStylePr>
    </w:style>
    """
    style_el = etree.fromstring(style_xml)
    doc.styles.element.append(style_el)


def main():
    if len(sys.argv) < 3:
        print("Usage: generate_docx.py <project_name> <output_path> [input_json_file]", file=sys.stderr)
        sys.exit(1)

    project_name = sys.argv[1]
    output_path = sys.argv[2]
    input_path = sys.argv[3] if len(sys.argv) > 3 else None

    if input_path:
        with open(input_path, "r") as f:
            sessions = json.load(f)
    else:
        sessions = json.loads(sys.stdin.read())

    doc = Document()

    # Set Garamond as default font
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(11)
    for heading_level in range(1, 4):
        hs = doc.styles[f"Heading {heading_level}"]
        hs.font.name = "Garamond"

    doc.add_heading(f"Arbeitsverlauf \u2013 LLM-Coding \u00d7 {project_name}", level=1)

    add_list_table_1_light_style(doc)
    table = doc.add_table(rows=1, cols=5, style="List Table 1 Light")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths via tblGrid (percentage-based, adapts to page)
    # Nr:4%, Datum:14%, Modell:10%, Umfang/Modus:19%, Commits:53%
    tbl_grid = table._tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        table._tbl.remove(tbl_grid)
    tbl_grid = etree.SubElement(table._tbl, qn("w:tblGrid"))
    grid_widths = [363, 1270, 907, 1724, 4808]  # in twips, total ~9072 (~6.3")
    for w in grid_widths:
        col = etree.SubElement(tbl_grid, qn("w:gridCol"))
        col.set(qn("w:w"), str(w))
    # Insert tblGrid right after tblPr
    tbl_pr = table._tbl.tblPr
    tbl_pr.addnext(tbl_grid)

    headers = ["Nr.", "Datum", "Modell", "Modus und Umfang", "Commits"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for s in sessions:
        if s.get("type") == "tombstones":
            continue  # rendered after regular entries
        row = table.add_row()
        row.cells[0].text = str(s["nr"])
        row.cells[1].text = s["datum"]
        row.cells[2].text = s.get("modell", "")
        row.cells[3].text = s["umfang"]

        cell = row.cells[4]
        cell.text = ""
        p = cell.paragraphs[0]
        topic = s.get("topic", "")
        if topic:
            run = p.add_run(topic)
            run.font.size = Pt(9)

        commits = s.get("commits", [])
        if commits:
            p_commits = cell.add_paragraph() if topic else p
            commit_text = ", ".join(commits)
            run = p_commits.add_run(commit_text)
            run.font.size = Pt(7)
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(128, 128, 128)

        for c in row.cells:
            for para in c.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

    # Render tombstone rows (if present)
    tombstone = next((s for s in sessions if s.get("type") == "tombstones"), None)
    if tombstone:
        entries = tombstone.get("entries", [])
        if entries:
            first_row_idx = len(table.rows)
            for entry in entries:
                row = table.add_row()
                row.cells[0].text = str(entry["nr"])
                row.cells[1].text = entry.get("datum", "")
                # Leave cells 2-3 empty (will be merged below)
                cell = row.cells[4]
                cell.text = ""
                p = cell.paragraphs[0]
                topic = entry.get("topic", "")
                if topic:
                    run = p.add_run(topic)
                    run.font.size = Pt(9)
                commits = entry.get("commits", [])
                if commits:
                    p_commits = cell.add_paragraph() if topic else p
                    run = p_commits.add_run(", ".join(commits))
                    run.font.size = Pt(7)
                    run.font.name = "Courier New"
                    run.font.color.rgb = RGBColor(128, 128, 128)
                for c in row.cells:
                    for para in c.paragraphs:
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)

            # Add top border to first tombstone row as visual divider
            W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            first_row = table.rows[first_row_idx]
            for cell in first_row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = etree.SubElement(tc_pr, qn("w:tcBorders"))
                top = etree.SubElement(borders, qn("w:top"))
                top.set(qn("w:val"), "single")
                top.set(qn("w:sz"), "4")
                top.set(qn("w:space"), "0")
                top.set(qn("w:color"), "666666")

            # Merge cells 2-3 across all tombstone rows (keep Nr. and Datum separate)
            last_row_idx = len(table.rows) - 1
            merged = table.cell(first_row_idx, 2).merge(table.cell(last_row_idx, 3))
            merged.text = ""
            p = merged.paragraphs[0]
            count = tombstone["count"]
            label = (
                f"{count} Sitzung{'en' if count != 1 else ''}\n"
                "(genauer Umfang unbekannt,\n"
                "Verlauf rekonstruiert\n"
                "aus Commits)"
            )
            run = p.add_run(label)
            run.font.size = Pt(8)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    # Date range from sessions
    dates = [s["datum"] for s in sessions if s.get("type") != "tombstones"]
    if not dates:
        # Tombstone-only: collect dates from tombstone entries
        for s in sessions:
            if s.get("type") == "tombstones":
                dates = [e["datum"] for e in s.get("entries", [])]
    date_range = f"{dates[0]} \u2013 {dates[-1]}" if len(dates) > 1 else (dates[0] if dates else "unbekannt")

    doc.add_paragraph("")
    entry_count = sum(1 for s in sessions if s.get("type") != "tombstones")
    tombstone_entries = sum(len(s.get("entries", [])) for s in sessions if s.get("type") == "tombstones")
    total_entries = entry_count + tombstone_entries
    footer = doc.add_paragraph(f"{total_entries} Eintr\u00e4ge | {date_range} | Projekt: {project_name}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(output_path)
    print(f"Saved to {output_path}")

    # Open the generated file
    subprocess.Popen(["open", output_path])


if __name__ == "__main__":
    main()
